from __future__ import annotations

import re
import shutil
import zipfile
from collections import Counter
from io import BytesIO
from pathlib import Path

from flask import Flask, jsonify, render_template, request


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 15 * 1024 * 1024

STOP_WORDS = {
    "about", "after", "also", "and", "are", "been", "being", "but", "can",
    "for", "from", "have", "into", "its", "job", "more", "our", "that",
    "the", "their", "this", "through", "using", "with", "will", "you", "your",
}


def words(text: str) -> list[str]:
    return [w.lower().strip(".") for w in re.findall(r"[A-Za-z][A-Za-z+#.]{2,}", text)]


def run_ocr(image) -> str:
    try:
        import pytesseract

        if not shutil.which("tesseract"):
            windows_tesseract = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
            if windows_tesseract.exists():
                pytesseract.pytesseract.tesseract_cmd = str(windows_tesseract)
            else:
                return ""
        return pytesseract.image_to_string(image, config="--oem 3 --psm 6")
    except Exception:
        return ""


def extract_pdf(raw: bytes) -> str:
    """Extract native PDF text, repair recoverable PDFs, then fall back to OCR if available."""
    extracted_texts = []

    # Method 1: pypdf (Pure Python, fast and reliable in serverless environments like Vercel)
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(raw), strict=False)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValueError("This PDF is password-protected. Remove the password and upload it again.") from exc

        pypdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if len(pypdf_text.strip()) >= 30:
            return pypdf_text

        # Try layout extraction mode for complex PDF formatting
        try:
            pypdf_layout = "\n".join(page.extract_text(extraction_mode="layout") or "" for page in reader.pages)
            if len(pypdf_layout.strip()) >= 30:
                return pypdf_layout
        except Exception:
            pass

        if pypdf_text.strip():
            extracted_texts.append(pypdf_text)
    except ValueError:
        raise
    except Exception:
        pass

    # Method 2: pdfplumber (Handles complex font maps & tables)
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(raw)) as pdf:
            plumber_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if len(plumber_text.strip()) >= 30:
                return plumber_text
            if plumber_text.strip():
                extracted_texts.append(plumber_text)
    except Exception:
        pass

    # Method 3: PyMuPDF / fitz
    document = None
    try:
        import fitz
        document = fitz.open(stream=raw, filetype="pdf")
        fitz_text = "\n".join(page.get_text("text") for page in document)
        if len(fitz_text.strip()) >= 30:
            return fitz_text
        if fitz_text.strip():
            extracted_texts.append(fitz_text)
    except Exception:
        pass

    # Method 4: Best partial extraction candidate
    for text in extracted_texts:
        if len(text.strip()) >= 30:
            return text

    # Method 5: OCR fallback if Tesseract is installed
    if document is not None:
        try:
            from PIL import Image
            ocr_pages = []
            for page in document:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                ocr = run_ocr(image)
                if ocr.strip():
                    ocr_pages.append(ocr)
            full_ocr = "\n".join(ocr_pages)
            if len(full_ocr.strip()) >= 30:
                return full_ocr
        except Exception:
            pass

    raise ValueError(
        "This PDF could not be read or contains scanned images without readable text. "
        "Please upload a text-based PDF or the original DOCX file."
    )


def extract_text(upload) -> str:
    suffix = Path(upload.filename or "").suffix.lower()
    raw = upload.read()
    if not raw:
        raise ValueError("The uploaded file is empty. Choose the resume again.")
    if suffix in {".txt", ".md"}:
        return raw.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        return extract_pdf(raw)
    if suffix == ".docx":
        try:
            from docx import Document
            document = Document(BytesIO(raw))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            for section in document.sections:
                parts.extend(p.text for p in section.header.paragraphs)
                parts.extend(p.text for p in section.footer.paragraphs)
            return "\n".join(parts)
        except ImportError:
            raise ValueError("DOCX support requires python-docx. Please install it.")
        except (zipfile.BadZipFile, ValueError) as exc:
            raise ValueError("This DOCX is damaged. Open it in Word and save a fresh copy.") from exc
        except Exception:
            raise ValueError("Could not read this DOCX file. Save it as a new DOCX or PDF and try again.")
    if suffix == ".rtf":
        text = raw.decode("utf-8", errors="ignore")
        return re.sub(r"\\[a-z]+\d* ?|[{}]", " ", text)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        try:
            from PIL import Image
            ocr_text = run_ocr(Image.open(BytesIO(raw)).convert("RGB"))
            if not ocr_text.strip():
                raise ValueError("Image OCR is not available on this server environment. Upload a text PDF or DOCX resume.")
            return ocr_text
        except ValueError:
            raise
        except Exception:
            raise ValueError("Image OCR is not available on this server environment. Upload a text PDF or DOCX resume.")
    if suffix == ".doc":
        raise ValueError("Legacy .DOC files are not reliably readable. Open it in Word and save as DOCX or PDF.")
    raise ValueError("Upload a PDF, DOCX, RTF, TXT, Markdown, PNG, JPG, WebP, BMP, or TIFF resume.")


@app.get("/")
def index():
    return render_template("index.html")


@app.get("/features")
def features():
    return render_template("features.html", page="features")


@app.get("/how-it-works")
def how_it_works():
    return render_template("how.html", page="how")


@app.get("/screen-resume")
def screening_page():
    return render_template("screen.html", page="screen")


@app.get("/why-us")
def why_us():
    return render_template("why.html", page="why")


@app.route("/contact", methods=["GET", "POST"])
def contact():
    sent = request.method == "POST"
    return render_template("contact.html", page="contact", sent=sent)


@app.errorhandler(413)
def upload_too_large(_error):
    return jsonify(error="This resume is larger than 15 MB. Compress it or upload a smaller file."), 413


@app.errorhandler(500)
def api_server_error(_error):
    if request.path.startswith("/api/"):
        return jsonify(error="The resume could not be processed. Try exporting it as a new PDF or DOCX."), 500
    return "Something went wrong.", 500


@app.post("/api/screen")
def screen_resume():
    upload = request.files.get("resume")
    description = request.form.get("job_description", "").strip()
    if not upload or not upload.filename:
        return jsonify(error="Choose a resume to continue."), 400
    if len(description) < 40:
        return jsonify(error="Add a more detailed job description (at least 40 characters)."), 400

    try:
        resume_text = extract_text(upload)
    except ValueError as exc:
        return jsonify(error=str(exc)), 400
    except Exception as exc:
        app.logger.exception("Resume extraction failed: %s", exc)
        return jsonify(error="This PDF could not be decoded. Try exporting it again or upload the original DOCX."), 422
    if len(resume_text.strip()) < 30:
        return jsonify(error="This resume contains too little readable text. If it is scanned, export it with OCR or upload the original DOCX."), 400

    job_terms = [w for w in words(description) if w not in STOP_WORDS]
    resume_terms = set(words(resume_text))
    ranked = [term for term, _ in Counter(job_terms).most_common(18)]
    matched = [term for term in ranked if term in resume_terms]
    missing = [term for term in ranked if term not in resume_terms]
    keyword_score = len(matched) / max(len(ranked), 1)

    experience_hits = len(re.findall(r"\b(?:19|20)\d{2}\b|\b\d+\+?\s+years?\b", resume_text, re.I))
    section_hits = sum(bool(re.search(rf"\b{s}\b", resume_text, re.I)) for s in ("experience", "education", "skills", "projects"))
    structure_score = min(1, section_hits / 3)
    evidence_score = min(1, (experience_hits + len(re.findall(r"\b\d+%\b", resume_text))) / 5)
    score = round(100 * (0.68 * keyword_score + 0.20 * structure_score + 0.12 * evidence_score))
    score = max(18, min(96, score))

    if score >= 75:
        verdict, summary = "Strong match", "The resume aligns well with the role and shows clear supporting evidence."
    elif score >= 52:
        verdict, summary = "Promising match", "There is meaningful overlap, with a few gaps worth validating in an interview."
    else:
        verdict, summary = "Developing match", "The profile has potential, but the resume does not yet demonstrate several role priorities."

    section_names = ("summary", "experience", "education", "skills", "projects")
    found_sections = [s.title() for s in section_names if re.search(rf"\b{s}\b", resume_text, re.I)]
    absent_sections = [s.title() for s in section_names if s.title() not in found_sections]
    has_email = bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", resume_text))
    has_phone = bool(re.search(r"(?:\+?\d[\s().-]*){8,}", resume_text))
    has_link = bool(re.search(r"linkedin\.com|github\.com|https?://", resume_text, re.I))
    metrics = len(re.findall(r"(?:\b\d+(?:\.\d+)?%|\$\s?\d+|\b\d+\+\s+(?:users|clients|projects|teams)\b)", resume_text, re.I))
    bullet_count = len(re.findall(r"(?m)^\s*(?:[-•▪●]|\d+[.)])\s+", resume_text))
    word_count = len(words(resume_text))
    contact_score = (8 if has_email else 0) + (7 if has_phone else 0) + (5 if has_link else 0)
    sections_score = min(30, len(found_sections) * 6)
    length_score = 15 if 300 <= word_count <= 900 else (10 if 180 <= word_count <= 1100 else 5)
    bullet_score = min(15, bullet_count * 2)
    metric_score = min(10, metrics * 3)
    format_score = 10 if Path(upload.filename or "").suffix.lower() in {".pdf", ".docx"} else 7
    ats_score = min(100, contact_score + sections_score + length_score + bullet_score + metric_score + format_score)
    ats_label = "ATS friendly" if ats_score >= 75 else ("Needs ATS improvements" if ats_score >= 50 else "Not ATS ready")

    strengths = []
    if matched: strengths.append(f"Matches {len(matched)} important terms from the job description, including {', '.join(matched[:4])}.")
    if metrics: strengths.append(f"Uses {metrics} measurable achievement signal{'s' if metrics != 1 else ''}, which strengthens credibility.")
    if len(found_sections) >= 4: strengths.append("Uses clear, recognizable resume sections that ATS systems can navigate.")
    if has_email and has_phone: strengths.append("Includes essential contact information for recruiter follow-up.")
    if not strengths: strengths.append("The resume is readable, but it needs clearer role-specific evidence to create a strong first impression.")

    recommendations = []
    if missing:
        recommendations.append({"priority": "High", "title": "Add missing role language", "detail": f"Where truthful, demonstrate these requirements inside experience bullets: {', '.join(missing[:6])}. Do not only list them in Skills."})
    if metrics < 3:
        recommendations.append({"priority": "High", "title": "Quantify your impact", "detail": "Rewrite duty-based bullets as achievements using scale, speed, revenue, reliability, users, or percentage improvement."})
    if "Summary" not in found_sections:
        recommendations.append({"priority": "Medium", "title": "Add a targeted professional summary", "detail": "Open with 2–3 lines naming your role, years of relevant experience, strongest domain, and one measurable result."})
    if absent_sections:
        recommendations.append({"priority": "Medium", "title": "Use standard ATS headings", "detail": f"Add or clearly label relevant sections: {', '.join(absent_sections)}. Avoid creative headings that parsers may miss."})
    if bullet_count < 5:
        recommendations.append({"priority": "Medium", "title": "Improve scanability", "detail": "Use 3–5 concise bullets per recent role. Begin with strong action verbs and keep each bullet focused on one outcome."})
    if not has_link:
        recommendations.append({"priority": "Low", "title": "Add a professional profile", "detail": "Include a clean LinkedIn or GitHub URL when it supports the role."})

    impression = [
        "Lead with the experience and technologies most relevant to this exact role.",
        "Use the pattern: action + task + technology + measurable result.",
        "Remove generic claims such as “hard-working”; prove them with outcomes.",
        "Keep formatting simple: one column, standard headings, consistent dates, and no text inside graphics.",
    ]

    return jsonify(
        score=score,
        verdict=verdict,
        summary=summary,
        matched=matched[:8],
        missing=missing[:6],
        stats={"keywords": len(matched), "sections": section_hits, "evidence": experience_hits},
        ats={"score": ats_score, "label": ats_label, "word_count": word_count, "bullets": bullet_count, "metrics": metrics, "contact": contact_score, "sections": sections_score, "format": format_score},
        sections={"found": found_sections, "missing": absent_sections},
        strengths=strengths[:4],
        recommendations=recommendations[:6],
        impression=impression,
    )


if __name__ == "__main__":
    app.run(debug=True)
